from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any


def parse_docx(path: Path, text_limit: int = 20000) -> dict[str, Any]:
    """Extract DOCX core properties and plain text (static only)."""
    try:
        from docx import Document
    except Exception as exc:
        return {
            "metadata": {"error": f"python-docx unavailable: {exc}"},
            "text": "",
            "tags": ["docx", "parse_error"],
            "reasons": ["docx_parser_unavailable"],
        }

    reasons: list[str] = []

    try:
        doc = Document(str(path))
        core = doc.core_properties
        metadata = {
            "author": str(core.author or ""),
            "title": str(core.title or ""),
            "subject": str(core.subject or ""),
            "created": str(core.created or ""),
            "modified": str(core.modified or ""),
            "paragraph_count": len(doc.paragraphs),
        }

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text).strip()
        text = text[:text_limit]

        try:
            with zipfile.ZipFile(path) as archive:
                entries = [item.filename.lower() for item in archive.infolist()]
            has_vba = any("vba" in entry or "vbaProject.bin".lower() in entry for entry in entries)
            has_embeddings = any("embeddings/" in entry for entry in entries)
            metadata["has_vba_indicator"] = has_vba
            metadata["has_embedded_objects"] = has_embeddings
            if has_vba:
                reasons.append("docx_vba_macro_indicator")
            if has_embeddings:
                reasons.append("docx_embedded_object_indicator")
        except Exception:
            pass

        return {
            "metadata": metadata,
            "text": text,
            "tags": ["docx"],
            "reasons": reasons,
        }
    except Exception as exc:
        return {
            "metadata": {"error": f"docx_parse_failed: {exc}"},
            "text": "",
            "tags": ["docx", "parse_error"],
            "reasons": ["docx_parse_failed"],
        }
