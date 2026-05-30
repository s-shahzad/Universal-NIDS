
from __future__ import annotations

import json
import re
import textwrap
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

DEFAULT_TITLE = "Universal Hybrid Network Intrusion Detection System for Live, Offline, and Multi-Format Threat Analysis"

REFS: list[dict[str, str]] = [
    {
        "id": "paxson1998bro",
        "label": "[1]",
        "citation": 'V. Paxson, "Bro: A System for Detecting Network Intruders in Real-Time," USENIX Security Symposium, 1998. URL: https://www.usenix.org/legacy/publications/library/proceedings/sec98/full_papers/paxson/paxson.pdf',
        "bib": "@inproceedings{paxson1998bro, author={Vern Paxson}, title={Bro: A System for Detecting Network Intruders in Real-Time}, booktitle={Proceedings of the 7th USENIX Security Symposium}, year={1998}, url={https://www.usenix.org/legacy/publications/library/proceedings/sec98/full_papers/paxson/paxson.pdf}}",
    },
    {
        "id": "scarfone2007idps",
        "label": "[2]",
        "citation": 'K. Scarfone and P. Mell, Guide to Intrusion Detection and Prevention Systems (IDPS), NIST SP 800-94, 2007. URL: https://csrc.nist.gov/pubs/sp/800/94/final',
        "bib": "@techreport{scarfone2007idps, author={Karen Scarfone and Peter Mell}, title={Guide to Intrusion Detection and Prevention Systems (IDPS)}, institution={National Institute of Standards and Technology}, number={SP 800-94}, year={2007}, url={https://csrc.nist.gov/pubs/sp/800/94/final}}",
    },
    {
        "id": "dreger2006dynamic",
        "label": "[3]",
        "citation": 'H. Dreger, A. Feldmann, V. Paxson, and R. Sommer, "Dynamic Application-Layer Protocol Analysis for Network Intrusion Detection," USENIX Security Symposium, 2006. URL: https://www.usenix.org/legacy/events/sec06/tech/full_papers/dreger/dreger.pdf',
        "bib": "@inproceedings{dreger2006dynamic, author={Holger Dreger and Anja Feldmann and Vern Paxson and Robin Sommer}, title={Dynamic Application-Layer Protocol Analysis for Network Intrusion Detection}, booktitle={Proceedings of the 15th USENIX Security Symposium}, year={2006}, url={https://www.usenix.org/legacy/events/sec06/tech/full_papers/dreger/dreger.pdf}}",
    },
    {
        "id": "tavallaee2009kdd",
        "label": "[4]",
        "citation": 'M. Tavallaee, E. Bagheri, W. Lu, and A. A. Ghorbani, "A Detailed Analysis of the KDD CUP 99 Data Set," IEEE CISDA, 2009. URL: https://www.ee.torontomu.ca/~bagheri/papers/cisda.pdf',
        "bib": "@inproceedings{tavallaee2009kdd, author={Mahbod Tavallaee and Ebrahim Bagheri and Wei Lu and Ali A. Ghorbani}, title={A Detailed Analysis of the KDD CUP 99 Data Set}, booktitle={Proceedings of the IEEE Symposium on Computational Intelligence for Security and Defense Applications}, year={2009}, url={https://www.ee.torontomu.ca/~bagheri/papers/cisda.pdf}}",
    },
    {
        "id": "sharafaldin2018cicids",
        "label": "[5]",
        "citation": 'I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, "Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization," ICISSP, 2018. URL: https://www.scitepress.org/Papers/2018/66398/66398.pdf',
        "bib": "@inproceedings{sharafaldin2018cicids, author={Iman Sharafaldin and Arash Habibi Lashkari and Ali A. Ghorbani}, title={Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization}, booktitle={Proceedings of the 4th International Conference on Information Systems Security and Privacy}, year={2018}, url={https://www.scitepress.org/Papers/2018/66398/66398.pdf}}",
    },
    {
        "id": "ring2019survey",
        "label": "[6]",
        "citation": 'M. Ring, S. Wunderlich, D. Scheuring, D. Landes, and A. Hotho, "A Survey of Network-based Intrusion Detection Data Sets," Computers & Security, 2019. URL: https://arxiv.org/abs/1903.02460',
        "bib": "@article{ring2019survey, author={Markus Ring and Sarah Wunderlich and Dominik Scheuring and Dieter Landes and Andreas Hotho}, title={A Survey of Network-based Intrusion Detection Data Sets}, journal={Computers & Security}, volume={86}, pages={147--167}, year={2019}, url={https://arxiv.org/abs/1903.02460}}",
    },
    {
        "id": "mirsky2018kitsune",
        "label": "[7]",
        "citation": 'Y. Mirsky, T. Doitshman, Y. Elovici, and A. Shabtai, "Kitsune: An Ensemble of Autoencoders for Online Network Intrusion Detection," NDSS, 2018. URL: https://www.ndss-symposium.org/ndss-paper/kitsune-an-ensemble-of-autoencoders-for-online-network-intrusion-detection/',
        "bib": "@inproceedings{mirsky2018kitsune, author={Yisroel Mirsky and Tomer Doitshman and Yuval Elovici and Asaf Shabtai}, title={Kitsune: An Ensemble of Autoencoders for Online Network Intrusion Detection}, booktitle={Network and Distributed System Security Symposium}, year={2018}, url={https://www.ndss-symposium.org/ndss-paper/kitsune-an-ensemble-of-autoencoders-for-online-network-intrusion-detection/}}",
    },
    {
        "id": "seth2024conceptdrift",
        "label": "[8]",
        "citation": 'S. Seth, "Concept Drift-Based Intrusion Detection Using Incremental Learning," The Computer Journal, 2024. URL: https://academic.oup.com/comjnl/article-abstract/67/7/2529/7618465',
        "bib": "@article{seth2024conceptdrift, author={Sugandh Seth}, title={Concept Drift-Based Intrusion Detection Using Incremental Learning}, journal={The Computer Journal}, volume={67}, number={7}, pages={2529--2548}, year={2024}, url={https://academic.oup.com/comjnl/article-abstract/67/7/2529/7618465}}",
    },
]


def _root(repo_root: Path | str | None = None) -> Path:
    return Path(repo_root).resolve() if repo_root is not None else Path(__file__).resolve().parents[2]


def _read_json(path: Path, default: Any) -> Any:
    return default if not path.exists() else json.loads(path.read_text(encoding="utf-8"))


def _write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def _metadata(root: Path) -> dict[str, str]:
    path = root / "thesis" / "thesis_metadata.json"
    data = {
        "project_title": DEFAULT_TITLE,
        "author": "Shaik",
        "institution": "Institution to be updated",
        "supervisor": "Supervisor to be updated",
        "project_version": f"research-snapshot-{datetime.now(timezone.utc).strftime('%Y.%m.%d')}",
    }
    if path.exists():
        data.update({k: str(v) for k, v in _read_json(path, {}).items() if v is not None})
    else:
        _write(path, json.dumps(data, indent=2))
    return data


def _requirements(root: Path) -> list[str]:
    path = root / "requirements.txt"
    if not path.exists():
        return []
    return [line.strip() for line in path.read_text(encoding="utf-8-sig").splitlines() if line.strip() and not line.startswith("#")]


def _artifact_formats(root: Path) -> list[str]:
    parser_dir = root / "src" / "NIDS" / "artifacts" / "parsers"
    mapping = {
        "csv_parser.py": "CSV",
        "docx_parser.py": "DOCX / Word documents",
        "exe_parser.py": "PE / EXE files",
        "html_parser.py": "HTML",
        "json_parser.py": "JSON",
        "pdf_parser.py": "PDF",
        "py_parser.py": "Python scripts",
        "xlsx_parser.py": "XLSX spreadsheets",
        "zip_parser.py": "ZIP archives",
    }
    return sorted([label for name, label in mapping.items() if (parser_dir / name).exists()])


def _state(root: Path) -> dict[str, Any]:
    metrics = _read_json(root / "reports" / "ml_metrics.json", {})
    evaluation = _read_json(root / "reports" / "ml_evaluation.json", {})
    ledger = _read_json(root / "NIDS_TestLab" / "reports" / "attack_test_ledger.json", {})
    overlap = list(ledger.get("concurrent_overlap_runs") or [])
    return {
        "root": root,
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "metadata": _metadata(root),
        "metrics": metrics,
        "evaluation": evaluation,
        "ledger": ledger,
        "requirements": _requirements(root),
        "formats": _artifact_formats(root),
        "feature_columns": list(metrics.get("feature_columns") or []),
        "completed_network": list(ledger.get("completed_network_detections") or []),
        "completed_static": list(ledger.get("completed_static_families") or []),
        "completed_os": list(ledger.get("completed_os_defense_cases") or []),
        "overlap": overlap,
        "overlap_pass": len([item for item in overlap if item.get("status") == "pass"]),
        "overlap_partial": len([item for item in overlap if item.get("status") == "partial"]),
        "remaining_attack": list(ledger.get("remaining_attack_families") or []),
        "remaining_os": list(ledger.get("remaining_os_defense_cases") or []),
    }


def _refs_markdown() -> str:
    return "\n".join(f"{item['label']} {item['citation']}" for item in REFS)


def _refs_bib() -> str:
    return "\n\n".join(item["bib"] for item in REFS)


def _normalize_md(text: str) -> str:
    return "\n".join(line[8:] if line.startswith("        ") else line for line in text.splitlines())

def _architecture_mermaid() -> str:
    return textwrap.dedent(
        """
        flowchart LR
            A[Live Capture] --> P[Parser and Feature Extraction]
            B[Offline PCAP Replay] --> P
            C[Suricata and Zeek Adapters] --> P
            D[Artifact Intake] --> E[Artifact Parsers and Analyzer]
            P --> F[Signature Engine]
            P --> G[Statistical Anomaly Engine]
            P --> H[ML Router]
            H --> I[Supervised Ensemble]
            H --> J[Unsupervised Hybrid Detector]
            F --> K[Fusion Engine]
            G --> K
            I --> K
            J --> K
            K --> L[Alert Suppression and Storage]
            L --> M[Dashboard, Charts, Reports]
            E --> L
        """
    ).strip()


def _workflow_mermaid() -> str:
    return textwrap.dedent(
        """
        flowchart TD
            S1[1. Data ingestion] --> S2[2. Parsing and normalization]
            S2 --> S3[3. Feature extraction]
            S3 --> S4[4. Signature and anomaly analysis]
            S4 --> S5[5. ML scoring]
            S5 --> S6[6. Fusion decision]
            S6 --> S7[7. Alert persistence]
            S7 --> S8[8. Visualization and thesis evidence]
        """
    ).strip()


def _architecture_doc(state: dict[str, Any]) -> str:
    features = ", ".join(state["feature_columns"]) if state["feature_columns"] else "packet, protocol, and flow statistics"
    return textwrap.dedent(
        f"""
        # Architecture Documentation

        Generated: {state['generated_at']}

        Universal NIDS is implemented as a layered platform inside `NIDS_Workspace`. Live capture, offline replay, adapter ingest, artifact intake, detection, storage, and visualization all converge into one repository-grounded workflow.

        ## Component Map

        - Ingest: `src/NIDS/ingest/live.py` and `src/NIDS/ingest/offline.py`
        - Parsing: `src/NIDS/pipeline/parser.py`
        - Feature extraction: `src/NIDS/pipeline/features.py`
        - Detection: `src/NIDS/detect/`
        - Storage: `src/NIDS/storage/`
        - Visualization: `src/NIDS/visuals/`
        - Artifact analysis: `src/NIDS/artifacts/`

        ## Feature Set

        The current supervised report exposes: {features}.

        ## Diagram Sources

        - `thesis/diagrams/system_architecture.mmd`
        - `thesis/diagrams/threat_workflow.mmd`

        ```mermaid
        {_architecture_mermaid()}
        ```

        ```mermaid
        {_workflow_mermaid()}
        ```
        """
    ).strip()


def _algorithms_doc(state: dict[str, Any]) -> str:
    algorithms = ", ".join(state["metrics"].get("algorithms") or []) or "random_forest, extra_trees, hist_gradient_boosting, xgboost"
    return textwrap.dedent(
        f"""
        # Algorithms Documentation

        Generated: {state['generated_at']}

        ## Active Runtime Stack

        - Signature rules for deterministic known-pattern detection
        - Statistical anomaly detection using thresholds, EWMA-style logic, z-score signals, and DNS burst logic
        - Supervised ensemble classification using {algorithms}
        - Optional unsupervised scoring using Isolation Forest and a shallow autoencoder
        - Fusion-based final decision logic

        The current production path is hybrid by design. Historical SVM and Decision Tree experiments remain part of the project history but are not the active runtime path.
        """
    ).strip()


def _experiments_doc(state: dict[str, Any]) -> str:
    return textwrap.dedent(
        f"""
        # Experimental Evaluation Documentation

        Generated: {state['generated_at']}

        ## Metrics

        - Training accuracy: {float(state['metrics'].get('accuracy', 0.0)):.5f}
        - Training weighted F1: {float(state['metrics'].get('f1_weighted', 0.0)):.5f}
        - Evaluation accuracy: {float(state['evaluation'].get('accuracy', 0.0)):.5f}
        - Evaluation weighted F1: {float(state['evaluation'].get('f1_weighted', 0.0)):.5f}

        ## Validation Summary

        - Completed network detections: {len(state['completed_network'])}
        - Completed static-malware family passes: {len(state['completed_static'])}
        - Completed OS-defense cases: {len(state['completed_os'])}
        - Overlap pass runs: {state['overlap_pass']}
        - Overlap partial runs: {state['overlap_partial']}
        - Remaining attack families: {len(state['remaining_attack'])}
        - Remaining OS cases: {len(state['remaining_os'])}
        """
    ).strip()


def _build_master_markdown(state: dict[str, Any]) -> str:
    meta = state["metadata"]
    metrics = state["metrics"]
    evaluation = state["evaluation"]
    ledger = state["ledger"]
    formats = ", ".join(state["formats"]) or "CSV, JSON, DOCX, HTML, PDF, EXE, Python, XLSX, ZIP"
    features = ", ".join(state["feature_columns"]) or "packet and flow statistics"
    requirements = ", ".join(state["requirements"]) or "the repository requirements set"
    network_cases = "\n".join(f"- {item['attack_case']}: {item['status']}" for item in state["completed_network"])
    static_cases = "\n".join(f"- {item['family']}: {item['status']}" for item in state["completed_static"])
    os_cases = "\n".join(f"- {item['attack_case']}: {item['status']}" for item in state["completed_os"])
    overlap_cases = "\n".join(f"- {item['run_name']}: {item['status']} - {item['finding']}" for item in state["overlap"][:6])
    remaining_attack = "\n".join(f"- {item}" for item in ledger.get("remaining_attack_families", []))
    remaining_os = "\n".join(f"- {item}" for item in ledger.get("remaining_os_defense_cases", []))
    os_case_count = len(state["completed_os"])
    os_case_phrase = "no completed OS-defense cases" if os_case_count == 0 else f"{os_case_count} completed OS-defense case{'s' if os_case_count != 1 else ''}"
    os_focus = state["remaining_os"][0] if state["remaining_os"] else "additional OS hardening validation"
    return textwrap.dedent(
        f"""
        # 1. Title Page

        ## Project Title
        {meta['project_title']}

        ## Author
        {meta['author']}

        ## Institution
        {meta['institution']}

        ## Supervisor
        {meta['supervisor']}

        ## Date
        {datetime.now(timezone.utc).strftime('%B %d, %Y')}

        ## Project Version
        {meta['project_version']}

        # 2. Abstract

        Universal NIDS addresses the practical gap between traditional packet-centric intrusion-detection systems and the heterogeneous evidence generated in modern security operations. The implemented system in `NIDS_Workspace` combines live monitoring, offline replay, artifact triage, supervised learning, unsupervised anomaly scoring, fusion-based alerting, visualization, and thesis-grade evidence retention. The runtime supports live interfaces, offline PCAPs, optional Suricata and Zeek adapters, and static handling of {formats}. Rather than relying on a single classifier, the current system combines signature logic, statistical anomaly detection, a supervised ensemble, optional unsupervised scoring, and a weighted fusion engine. On the current benchmark-derived labeled flow corpus of {metrics.get('samples_total', 'unknown')} samples, the supervised path achieved training accuracy {float(metrics.get('accuracy', 0.0)):.5f} and weighted F1 {float(metrics.get('f1_weighted', 0.0)):.5f}; the recorded evaluation run achieved accuracy {float(evaluation.get('accuracy', 0.0)):.5f} and weighted F1 {float(evaluation.get('f1_weighted', 0.0)):.5f}. Experimental evidence in the isolated VirtualBox lab confirms detection of web-path abuse, DNS bursts, brute-force activity, port scanning, DoS-rate anomalies, and {os_case_phrase}. The main research contribution is therefore a real repository-grounded hybrid NIDS and a continuously updated thesis record that preserves both successful detections and unresolved gaps.

        # 3. Introduction

        ## 3.1 Background
        Network intrusion detection is still a central cybersecurity problem because defenders must classify legitimate and malicious activity under high traffic volume, protocol diversity, and incomplete labels. Foundational work established the value of real-time protocol-aware monitoring and operational IDS design [1]-[3].

        ## 3.2 Motivation
        Modern security work rarely involves packets alone. Analysts must correlate PCAPs, logs, documents, scripts, executables, and model outputs. This repository was built to keep those evidence paths inside one system and one research narrative.

        ## 3.3 Need for an Advanced NIDS
        The literature shows that fixed signatures alone are insufficient, while pure anomaly models and purely supervised models also have clear weaknesses [4]-[8]. A hybrid architecture is therefore justified for this project.

        ## 3.4 Research Objective Statement
        The research objective is to build and continuously document a universal hybrid NIDS that supports heterogeneous data, live and offline analysis, machine-learning-assisted detection, and reproducible attack evidence.

        # 4. Problem Statement

        ## 4.1 Heterogeneous Inputs
        Problem: security evidence arrives in many formats.
        Impact: analysts lose context when network and file evidence are separated.
        Existing Limitation: many IDS workflows only inspect traffic or only inspect files.
        Need for Improvement: one architecture should process both streams.

        ## 4.2 Real-Time Detection Constraints
        Problem: live scoring must remain fast enough for operational use.
        Impact: expensive scoring can degrade alert quality under load.
        Existing Limitation: single-engine designs are either too narrow or too fragile.
        Need for Improvement: hybrid detection with selective scoring is required.

        ## 4.3 Labeled-Data Dependence
        Problem: benchmark-driven accuracy can overstate deployment readiness.
        Impact: models may not transfer cleanly to live environments.
        Existing Limitation: many IDS claims rely on data with realism limitations.
        Need for Improvement: labeled and unlabeled paths must coexist.

        ## 4.4 Research Traceability
        Problem: security experiments often lose failed runs.
        Impact: reviewers cannot verify what was fixed and what remains open.
        Existing Limitation: engineering notes, metrics, and screenshots are usually fragmented.
        Need for Improvement: the thesis record must evolve with the repository.

        # 5. Research Objectives

        1. Build a universal NIDS for live and offline monitoring.
        2. Support heterogeneous file and network evidence.
        3. Implement hybrid detection instead of a single-model pipeline.
        4. Preserve an evidence trail for attack testing, malware triage, and OS-defense validation.
        5. Provide visualization, reporting, and reproducible research outputs.
        6. Keep the architecture modular and scalable.

        # 6. Literature Review

        ## 6.1 Signature and Protocol-Aware IDS
        Early work by Paxson and later operational guidance from NIST demonstrate why protocol-aware inspection is foundational for practical intrusion detection [1][2]. Dreger et al. further showed that application-layer analysis must not rely blindly on destination ports [3].

        ## 6.2 Dataset Realism and Machine Learning
        Tavallaee et al., Sharafaldin et al., and Ring et al. collectively show that data quality, traffic realism, and benchmark diversity strongly influence IDS claims [4][5][6]. This is directly relevant because the current supervised model still depends on benchmark-derived training data.

        ## 6.3 Online Anomaly Detection and Representation Learning
        Kitsune is relevant because it demonstrates that lightweight representation learning can be used for online intrusion detection [7]. The present repository takes a pragmatic hybrid route rather than claiming a full deep-learning runtime.

        ## 6.4 Concept Drift
        Recent work on concept drift indicates that intrusion-detection models must adapt to changing traffic distributions [8]. This remains future work in the current repository.

        ## 6.5 Research Inference
        Taken together, these sources imply that no single detector family is sufficient. This thesis therefore adopts a hybrid design as an explicit inference from the literature.

        # 7. System Architecture

        ## 7.1 System Layers
        The system is organized into ingestion, parsing, feature extraction, detection, storage, and visualization layers.

        ## 7.2 Offline Pipeline
        Offline PCAP replay enters through `src/NIDS/ingest/offline.py`, then follows the same parsing and detection path as live traffic.

        ## 7.3 Live Monitoring Pipeline
        Live capture enters through `src/NIDS/ingest/live.py`, is normalized by the pipeline, scored by multiple engines, and persisted to SQLite and JSONL outputs.

        ## 7.4 Component Interaction
        Runtime coordination is centered in `src/NIDS/runtime.py`, with detection modules in `src/NIDS/detect/`, storage in `src/NIDS/storage/`, and visualization in `src/NIDS/visuals/`.

        ## 7.5 Diagram Sources
        - `thesis/diagrams/system_architecture.mmd`
        - `thesis/diagrams/threat_workflow.mmd`

        # 8. Data Handling and Classification

        ## 8.1 Multi-Format Support
        The system currently supports {formats}, plus logs, live packets, offline PCAPs, and adapter JSON feeds.

        ## 8.2 Parsing and Normalization
        Artifact parsers operate in `src/NIDS/artifacts/parsers/`, while network normalization is handled by `src/NIDS/pipeline/parser.py`.

        ## 8.3 Feature Extraction
        The current feature set includes {features}.

        ## 8.4 Labeled and Unlabeled Paths
        Labeled flow data supports supervised training and evaluation, while unlabeled traffic can be scored through the unsupervised path with warmup calibration and baseline persistence.

        # 9. Detection Algorithms

        ## 9.1 Signature Rules
        YAML-driven signature rules provide deterministic detection of known malicious patterns.

        ## 9.2 Statistical Anomaly Detection
        Threshold logic, EWMA-style behavior, z-score logic, and DNS burst handling provide fast interpretable streaming detection.

        ## 9.3 Supervised Ensemble
        The current active supervised ensemble uses {', '.join(metrics.get('algorithms') or []) or 'random_forest, extra_trees, hist_gradient_boosting, xgboost'}.

        ## 9.4 Unsupervised Hybrid Detection
        The optional unsupervised path combines Isolation Forest with a shallow autoencoder.

        ## 9.5 Fusion
        The fusion engine combines signature, anomaly, supervised, and unsupervised signals into one final decision.

        # 10. Threat Detection Workflow

        1. Ingest data from a live interface, offline PCAP, or adapter source.
        2. Parse and normalize the event.
        3. Extract features.
        4. Evaluate signature and anomaly rules.
        5. Score the event with supervised and optional unsupervised ML.
        6. Fuse the active signals.
        7. Persist alerts, flows, and metrics.
        8. Expose the result to reports, charts, dashboards, and thesis evidence.

        # 11. Graphical Threat Pattern Analysis

        ## 11.1 Offline Analytics
        `src/NIDS/visuals/charts.py` and `src/NIDS/visuals/export.py` support chart generation and evidence packaging.

        ## 11.2 Live Dashboard
        `src/NIDS/visuals/dashboard.py` exposes rolling metrics, alerts, sensor comparison, and incident-oriented views.

        ## 11.3 Analyst Value
        Visualization improves interpretability by showing how multiple detectors align across time and severity.

        # 12. Implementation Details

        ## 12.1 Languages and Core Modules
        The project is primarily implemented in Python under `src/NIDS/`.

        ## 12.2 Frameworks and Libraries
        The active dependency set includes {requirements}.

        ## 12.3 Storage and Reporting
        SQLite and JSONL are the primary runtime outputs, while Markdown and DOCX support formal reporting.

        ## 12.4 Tooling
        Testing is performed with `pytest`; visualization relies on Plotly and Kaleido; dashboard services use FastAPI and Uvicorn.

        # 13. Experimental Evaluation

        ## 13.1 Datasets and Evidence Sources
        The current supervised path is still trained on benchmark-derived data, while experimental validation also uses offline replay, live VM lab scenarios, and static artifact families.

        ## 13.2 Metrics
        The recorded metrics include accuracy, weighted precision, weighted recall, weighted F1, and confusion matrices.

        ## 13.3 Performance Summary
        - Training samples: {metrics.get('samples_total', 'unknown')}
        - Training accuracy: {float(metrics.get('accuracy', 0.0)):.5f}
        - Training weighted F1: {float(metrics.get('f1_weighted', 0.0)):.5f}
        - Evaluation samples: {evaluation.get('samples', 'unknown')}
        - Evaluation accuracy: {float(evaluation.get('accuracy', 0.0)):.5f}
        - Evaluation weighted F1: {float(evaluation.get('f1_weighted', 0.0)):.5f}

        ## 13.4 Validated Cases
        {network_cases or '- No network validation entries recorded.'}

        ## 13.5 Static and OS Evidence
        {static_cases or '- No static triage entries recorded.'}
        {os_cases or '- No OS-defense entries recorded.'}

        ## 13.6 Concurrent Overlap Evidence
        {overlap_cases or '- No overlap entries recorded.'}

        # 14. Improvements Over Existing Systems

        ## 14.1 Multi-Format Ingestion
        Problem: traditional IDS platforms rarely unify packets and suspicious files.
        Previous Limitation: analysts must switch tools.
        Proposed Solution: implemented artifact intake plus network pipelines.
        Result: one workspace now preserves both contexts.

        ## 14.2 Hybrid Detection
        Problem: no single detector family is sufficient.
        Previous Limitation: rule-only and model-only designs both miss important behaviors.
        Proposed Solution: signature, anomaly, supervised, unsupervised, and fusion logic are combined.
        Result: the runtime can preserve deterministic alerts while still handling unknown patterns.

        ## 14.3 Evidence Preservation
        Problem: failed detections are often lost.
        Previous Limitation: security projects over-report only successful cases.
        Proposed Solution: coverage matrix, test ledger, OS-defense plan, and thesis outputs are maintained together.
        Result: the current project keeps both passes and misses in scope.

        # 15. Current System Version

        ## 15.1 Version
        {meta['project_version']}

        ## 15.2 Completed Modules
        - Live and offline runtime
        - Signature, anomaly, supervised, unsupervised, and fusion detection
        - SQLite and JSONL persistence
        - Static artifact triage
        - Dashboard and chart export
        - Attack-lab evidence ledger and thesis documentation

        ## 15.3 Modules Under Development
        - Lower-rate periodic beacon detection and C2-oriented beacon analysis beyond the current burst-style DNS validation
        - More same-window multi-attack validation
        - Real-sample malware execution only inside the isolated lab
        - Next OS-defense focus: {os_focus}

        # 16. Future Work

        ## 16.1 Remaining Attack Families
        {remaining_attack or '- No remaining attack items listed.'}

        ## 16.2 Remaining OS Cases
        {remaining_os or '- No remaining OS items listed.'}

        ## 16.3 Research Expansion
        Future work should add drift handling, adversarial evaluation, SIEM integration, distributed deployment, and deeper sequence-aware models only where the data justifies them.

        # 17. Conclusion

        Universal NIDS is no longer just a detection script; it is a repository-grounded hybrid research platform. Its current value lies in the combination of operational code, measurable evaluation, isolated-lab validation, and honest preservation of unresolved gaps.

        # 18. References

        {_refs_markdown()}
        """
    ).strip() + "\n"


def _escape_tex(text: str) -> str:
    mapping = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
    }
    return "".join(mapping.get(ch, ch) for ch in text)


def _build_thesis_tex(state: dict[str, Any]) -> str:
    meta = state["metadata"]
    metrics = state["metrics"]
    evaluation = state["evaluation"]
    reqs = ", ".join(state["requirements"]) or "the current repository requirements set"
    return textwrap.dedent(
        rf"""
        \documentclass[12pt]{{report}}
        \usepackage[margin=1in]{{geometry}}
        \usepackage{{hyperref}}
        \title{{{_escape_tex(meta['project_title'])}}}
        \author{{{_escape_tex(meta['author'])}}}
        \date{{{_escape_tex(datetime.now(timezone.utc).strftime('%B %d, %Y'))}}}
        \begin{{document}}
        \begin{{titlepage}}
        \centering
        {{\LARGE {_escape_tex(meta['project_title'])} \par}}
        \vspace{{1cm}}
        {{\large Author: {_escape_tex(meta['author'])} \par}}
        {{\large Institution: {_escape_tex(meta['institution'])} \par}}
        {{\large Supervisor: {_escape_tex(meta['supervisor'])} \par}}
        {{\large Project Version: {_escape_tex(meta['project_version'])} \par}}
        \end{{titlepage}}
        \chapter*{{Abstract}}
        Universal NIDS is a repository-grounded hybrid intrusion-detection platform for live traffic, offline replay, artifact triage, machine learning, fusion-based alerting, and research-quality evidence retention.
        \chapter{{Introduction}}
        The project addresses heterogeneous evidence, operational constraints, and the need for a continuously updated research record.
        \chapter{{Literature Review}}
        The thesis draws on protocol-aware IDS, operational guidance, dataset-evaluation studies, online anomaly detection, and concept-drift literature \cite{{paxson1998bro,scarfone2007idps,dreger2006dynamic,tavallaee2009kdd,sharafaldin2018cicids,ring2019survey,mirsky2018kitsune,seth2024conceptdrift}}.
        \chapter{{System and Evaluation Summary}}
        Current training accuracy is {_escape_tex(f"{float(metrics.get('accuracy', 0.0)):.5f}")} with weighted F1 {_escape_tex(f"{float(metrics.get('f1_weighted', 0.0)):.5f}")}. Current evaluation accuracy is {_escape_tex(f"{float(evaluation.get('accuracy', 0.0)):.5f}")} with weighted F1 {_escape_tex(f"{float(evaluation.get('f1_weighted', 0.0)):.5f}")}. The implementation uses {_escape_tex(reqs)}.
        \chapter{{Conclusion}}
        Universal NIDS provides a practical hybrid architecture and a research-grade evidence trail while openly retaining unresolved gaps for future work.
        \bibliographystyle{{IEEEtran}}
        \bibliography{{references}}
        \end{{document}}
        """
    ).strip() + "\n"


def _thesis_readme() -> str:
    return textwrap.dedent(
        """
        # Thesis Documentation Workflow

        Run from the repository root with:

        ```bash
        python -m nids thesis-docs
        ```

        Outputs stay under `NIDS_Workspace`:

        - `thesis/nids_thesis.tex`
        - `thesis/references.bib`
        - `thesis/diagrams/system_architecture.mmd`
        - `thesis/diagrams/threat_workflow.mmd`
        - `documentation/architecture.md`
        - `documentation/algorithms.md`
        - `documentation/experiments.md`
        - `NIDS_TestLab/reports/nids_project_master_thesis.md`
        - `NIDS_TestLab/reports/nids_project_master_thesis.docx`
        - `thesis/thesis_metadata.json`
        """
    ).strip()


def _render_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    first_h1 = True
    in_code = False
    code_lines: list[str] = []
    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            doc.add_heading(line[2:].strip(), level=0)
            first_h1 = False
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_thesis_documents(repo_root: Path | str | None = None, out_md: Path | str | None = None, out_docx: Path | str | None = None) -> dict[str, str]:
    root = _root(repo_root)
    state = _state(root)
    thesis_dir = root / "thesis"
    docs_dir = root / "documentation"
    diagrams_dir = thesis_dir / "diagrams"
    figures_dir = thesis_dir / "figures"
    reports_dir = root / "NIDS_TestLab" / "reports"
    md_text = _normalize_md(_build_master_markdown(state))
    md_path = Path(out_md).resolve() if out_md else reports_dir / "nids_project_master_thesis.md"
    docx_path = Path(out_docx).resolve() if out_docx else reports_dir / "nids_project_master_thesis.docx"
    _write(docs_dir / "architecture.md", _normalize_md(_architecture_doc(state)))
    _write(docs_dir / "algorithms.md", _normalize_md(_algorithms_doc(state)))
    _write(docs_dir / "experiments.md", _normalize_md(_experiments_doc(state)))
    _write(thesis_dir / "nids_thesis.tex", _build_thesis_tex(state))
    _write(thesis_dir / "references.bib", _refs_bib())
    _write(thesis_dir / "README.md", _normalize_md(_thesis_readme()))
    _write(diagrams_dir / "system_architecture.mmd", _architecture_mermaid())
    _write(diagrams_dir / "threat_workflow.mmd", _workflow_mermaid())
    _write(figures_dir / "README.md", "Rendered thesis figures can be stored here.\n")
    _write(md_path, md_text)
    _render_docx(md_text, docx_path)
    return {
        "markdown": str(md_path),
        "docx": str(docx_path),
        "thesis_tex": str(thesis_dir / "nids_thesis.tex"),
        "references_bib": str(thesis_dir / "references.bib"),
        "architecture_md": str(docs_dir / "architecture.md"),
        "algorithms_md": str(docs_dir / "algorithms.md"),
        "experiments_md": str(docs_dir / "experiments.md"),
    }
